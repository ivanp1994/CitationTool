# -*- coding: utf-8 -*-
"""
Created on Tue Aug 13 08:18:48 2024

@author: ivanp
"""
import sys
import os
import time
import re
from functools import lru_cache

from typing import List, Dict, Optional, Tuple
from docx import Document
from docx.text.paragraph import Paragraph
import requests

import PyQt5.QtWidgets as qtw
from PyQt5.QtCore import Qt  # pylint: disable=E0611
# the above is false positive - see https://stackoverflow.com/questions/56726580/no-name-qapplication-in-module-pyqt5-qtwidgets-error-in-pylint




#
LREF_PATTERN = re.compile(r'\[LR[^\]]*\]')

CN_BASE_URL = "https://doi.org"
HEADERS = {"User-Agent": "python-requests/" + requests.__version__,
           "X-USER-AGENT": "python-requests/" + requests.__version__,
           "Accept": "text/x-bibliography; style = apa; locale = en-Us"}

#
@lru_cache(maxsize=256)
def fetch_request(url: str) -> Optional[str]:
    """
    Interface to fetch the request
    """
    r = requests.get(url, headers=HEADERS, allow_redirects=True)  # pylint: disable=C0103
    if not r.ok:
        return None
    r.encoding = "UTF-8"
    return r.text


def extract_references(docx_path: str,
                       pattern: Optional[re.Pattern] = None) -> Dict[str, List[str]]:
    """
    Extracts all the references found in
    a document path that match the given pattern.
    """
    if pattern is None:
        pattern = LREF_PATTERN

    doc = Document(docx_path)
    doi_list = list()
    for para in doc.paragraphs:
        matches = pattern.findall(para.text)
        doi_list.extend(matches)

    reference_dict = {k: k.strip("]").strip("[LR:").split(";") for k in doi_list}
    return reference_dict


def merge_dois(doi_dict: Dict[str, List[str]]) -> List[str]:
    """
    Merges all DOIs from the extracted references into a single list.
    """
    return list({item for sublist in doi_dict.values() for item in sublist})


def intext_cit(literature_citation: str) -> str:
    """
    Creates an APA style intext citation from a given string

    """
    authors, year = literature_citation.split("(")[:2]
    year = year.split(")")[0]

    author_count = authors.count(".,") + authors.count("&")

    first_author = authors.split(",")[0].strip()
    if author_count == 2:
        second_author = " & " + \
            authors.split("&")[1].split(",")[0].strip()+", "
    # ., counts all authors and & counts last author
    # when there is one author this can be only 0
    elif author_count == 0:
        first_author = first_author + ","
        second_author = " "
    else:
        second_author = " et al., "

    intext_citation = first_author + second_author + year
    return intext_citation


def generate_replacer(lrtext_doi: Dict[str, List[str]],
                      doi_intext: Dict[str, str]) -> Tuple[Dict[str, str], Dict[str, str]]:
    """
    "lrtext_doi" is how references are in the raw file, e.g.
    "doi_intext" is how a citation is written during text, e.g.
    '(Zhang et al., 2024)'

    This returns two dictionaries - good replacer and bad replacer
    """
    good_replacer = dict()
    bad_replacer = dict()
    for intext, doi_list in lrtext_doi.items():
        replacement = "("
        one_citation = len(doi_list) == 1

        if one_citation:
            joiner = ")"
            end = ""
        else:
            joiner = "; "
            end = ")"

        for doi in doi_list:
            intext_doi = doi_intext.get(doi, None)
            if intext_doi is None:
                intext_doi = "BAD_DOI:"+doi
            replacement = replacement + intext_doi + joiner

        replacement = replacement.rstrip("; ") + end

        if "BAD_DOI" in replacement:
            bad_replacer[intext] = replacement
        else:
            good_replacer[intext] = replacement
    return good_replacer, bad_replacer


def replace_text_in_runs(paragraph: Paragraph,
                         replacements: Dict[str, str]) -> None:
    """
    Replaces a pattern in text
    Doesnt save comments
    """
    for key, value in replacements.items():

        if key in paragraph.text:
            full_text = ''.join(run.text for run in paragraph.runs)
            full_text = full_text.replace(key, value)

            start_index = 0
            for run in paragraph.runs:
                run_length = len(run.text)
                run.text = full_text[start_index:start_index + run_length]
                start_index += run_length


def save_document(doc_path: str, doc_out: str, replacement: Dict[str, str]) -> None:
    """
    Replaces a pattern in doc and then saves it as the output
    """
    doc = Document(doc_path)
    for paragraph in doc.paragraphs:
        replace_text_in_runs(paragraph, replacement)
    doc.save(doc_out)


def add_suffix(file_path: str, suffix: str) -> str:
    """
    Add a suffix to the filename in the given file path.
    """
    directory, filename = os.path.split(file_path)
    name, ext = os.path.splitext(filename)

    new_filename = f"{name}{suffix}{ext}"

    # Construct the new file path
    new_file_path = os.path.join(directory, new_filename)

    return new_file_path


def highlight_text(text: str, pattern: str, color_code: str) -> str:
    """Highlight occurrences of `pattern` in `text` with the specified `color_code`."""
    parts = text.split(pattern)
    highlighted_text = ''

    for i in range(len(parts) - 1):
        highlighted_text += parts[i]

        highlighted_text += f'<font color="{color_code}">{pattern}</font>'

    highlighted_text += parts[-1]

    return highlighted_text


def split_into_sentences(text: str) -> List[str]:
    """Split text into sentences using basic punctuation rules."""

    sentence_endings = re.compile(r"(?<=[.!?])\s+(?=[A-Z])")
    sentences = sentence_endings.split(text)
    # edge cases
    return [sentence.strip() for sentence in sentences if sentence.strip()]


def recognize_bad_dois(doc_path: str, patterns: List[str]) -> List[str]:
    """
    Recognizes bad doi and returns sentences containing them,
    with highlighting them
    """
    doc = Document(doc_path)

    # I think orange is nice but red is even bettern
    highlight_color = "red"
    bad_sentences = list()

    for para in doc.paragraphs:
        text = para.text
        sentences = split_into_sentences(text)
        for sentence in sentences:
            if any(pattern in sentence for pattern in patterns):
                highlighted_sentence = sentence
                for pattern in patterns:
                    if pattern in highlighted_sentence:
                        highlighted_sentence = highlight_text(
                            highlighted_sentence, pattern, highlight_color)
                        bad_sentences.append(highlighted_sentence)
    return bad_sentences


def save_literature_to_docx(literature_entries: List[str], last_path: str) -> str:
    """
    saves literature to document provided
    """
    # literature saving loging
    end_literature = sorted(list(literature_entries))
    _directory, _ = os.path.split(last_path)  # pylint: disable=W0631
    _literature = os.path.join(_directory, "Literature.docx")

    doc = Document()
    doc.add_heading('Literature References', level=1)
    for entry in end_literature:
        doc.add_paragraph(entry.strip())  # Strip trailing newline characters
    doc.save(_literature)
    return _literature


#%% GUI 
class FetchDOIApp(qtw.QMainWindow):
    """
    Simple fetcher qt Window
    """

    def __init__(self):
        super().__init__()
        self.lref_dict = {}  # {file : {"[LR...]":[doi,doi]}}
        self.doi_dict = {}  # doi : citation
        self._cit_intext = {}  # citation : intext
        self.file_paths = []  # files

        self.init_UI()

    def init_UI(self):  # pylint: disable=invalid-name
        "Intializes user interface"
        # Create main layout
        layout = qtw.QVBoxLayout()

        # Create QLabel for status messages
        self.statusLabel = qtw.QLabel('Select files to process...', self)  # pylint: disable=invalid-name
        self.statusLabel.setWordWrap(True)  # Allows text to wrap in QLabel

        # Create QScrollArea and configure
        self.scrollArea = qtw.QScrollArea(self)  # pylint: disable=invalid-name
        self.scrollArea.setWidgetResizable(True)  # Make the widget inside scrollArea resizable
        self.scrollArea.setFrameShape(qtw.QFrame.NoFrame)  # Remove border

        # Create container widget for QLabel
        self.statusContainer = qtw.QWidget()  # pylint: disable=invalid-name
        status_layout = qtw.QVBoxLayout()
        status_layout.addWidget(self.statusLabel)
        self.statusContainer.setLayout(status_layout)

        # Set container widget for the scroll area
        self.scrollArea.setWidget(self.statusContainer)

        # Create QProgressBar and configure
        self.progressBar = qtw.QProgressBar(self)  # pylint: disable=invalid-name
        self.progressBar.setAlignment(Qt.AlignCenter)

        # Add widgets to the main layout
        layout.addWidget(self.scrollArea)  # Expandable area for text
        layout.addWidget(self.progressBar)  # Positioned at the bottom

        # Create a central widget for the main window
        centralWidget = qtw.QWidget()  # pylint: disable=invalid-name
        centralWidget.setLayout(layout)
        self.setCentralWidget(centralWidget)

        # Configure the main window
        self.setWindowTitle('DOI Processor')
        self.setGeometry(100, 100, 400, 300)
        self.show()

        # Call the method to process files
        self.process_files()

    def process_files(self):
        """
        Handle opening files
        """
        # Open file dialog to select files
        file_dialog = qtw.QFileDialog(self)
        file_dialog.setFileMode(qtw.QFileDialog.ExistingFiles)
        file_dialog.setNameFilter("Word Documents (*.docx)")
        file_dialog.setAcceptMode(qtw.QFileDialog.AcceptOpen)
        file_dialog.setOption(qtw.QFileDialog.DontUseNativeDialog, True)
        if file_dialog.exec_():
            self.file_paths = file_dialog.selectedFiles()
        self.extract_and_fetch_dois()


    def end_session(self):
        " End if nothing is selected"
        # Close the main window and quit the application
        self.close()
        qtw.QApplication.quit()
        sys.exit(0)

    def extract_and_fetch_dois(self):
        "fetching dois from the DOCX"

        # Initialize progress bar
        total_files = len(self.file_paths)
        if total_files == 0:
            self.end_session()

        self.progressBar.setMaximum(total_files)
        self.progressBar.setValue(0)

        all_dois = []
        for idx, file_path in enumerate(self.file_paths):
            reference_dict = extract_references(file_path)
            self.lref_dict[file_path] = reference_dict
            all_dois.extend(merge_dois(reference_dict))
            self.progressBar.setValue(idx+1)
        all_dois = list(set(all_dois))
        _text = f"Extracted all DOIs - total of {len(all_dois)} DOIS"
        self.statusLabel.setText(_text)
        # Remove duplicates
        self.fetch_dois(list(set(all_dois)))

    def fetch_dois(self, all_dois: List[str]):
        """
        Fetches DOIs and updates the status label and progress bar.
        This is the time consuming part
        """

        time.sleep(0.5)
        self.progressBar.setValue(0)
        self.progressBar.setMaximum(len(all_dois))
        # Fetch DOIs
        doi_dictionary = {}
        previous_citation = "-1"
        for idx, ids in enumerate(all_dois):

            try:
                next_doi = f"\n..Fetching doi {all_dois[idx+1]}"
            except IndexError:
                next_doi = ""

            if previous_citation == "-1":
                valid_previous_doi = ""
            else:
                if previous_citation is None:
                    valid_previous_doi = "\nDOI was invalid"
                else:
                    valid_previous_doi = "\n" + previous_citation

            self.statusLabel.setText(f"Fetched doi {ids}\n{valid_previous_doi}{next_doi}")
            previous_citation = fetch_request(CN_BASE_URL + "/" + ids.strip())
            doi_dictionary[ids] = previous_citation
            time.sleep(0.004)
            self.progressBar.setValue(idx + 1)  # Update progress bar
            qtw.QApplication.processEvents()  # Process events to keep GUI responsive

        # Update status label with completion message
        self.statusLabel.setText('Done fetching DOIs!')
        self.progressBar.setValue(len(all_dois))  # Ensure progress bar reaches 100%
        time.sleep(1)
        self.doi_dict = doi_dictionary

        self.end_and_replace()

    def end_and_replace(self):
        """
        Ends the operation
        """
        self.statusLabel.setText("Replacing literature..")
        self.progressBar.setValue(0)
        self.progressBar.setMaximum(len(self.lref_dict))
        # time.sleep(5)

        _end_literature = sorted(list({v for v in self.doi_dict.values() if v is not None}))
        _lit_loc = save_literature_to_docx(_end_literature, self.file_paths[-1])
        self.statusLabel.setText(f'Saved literature to <font color = "green">{_lit_loc}</font>')
        qtw.QApplication.processEvents()  # Process events to keep GUI responsive

        time.sleep(5)

        # Now from doi dict which is DOI citation I want to create citation_intext
        self._cit_intext = {k: intext_cit(v)
                            for k, v in self.doi_dict.items() if v is not None}

        # {file : {"[LR...]":[doi,doi]}}
        good_replacer = {}
        bad_replacer = {}
        bad_sentences = ""
        display_text = ""
        idx = 0
        for file, wtdict in self.lref_dict.items():
            good, bad = generate_replacer(wtdict, self._cit_intext)
            good_replacer[file] = good
            bad_replacer[file] = bad

            file_name = os.path.basename(file)
            result_name = add_suffix(file, "_proc")

            save_document(file, result_name, good)

            idx = idx + 1
            self.progressBar.setValue(idx + 1)  # Update progress bar
            display_text = display_text + f'Replaced {len(good)} citation(s) in <font color="green">{file_name}</font> - {len(bad)} was/were bad<br>'
            self.statusLabel.setText(display_text)

            qtw.QApplication.processEvents()  # Process events to keep GUI responsive

            _bad_sentences = [f"\t{x}" for x in recognize_bad_dois(file, list(bad.keys()))]

            bad_sentences = bad_sentences + f'<font color="green">{file_name}</font><br>---<br>' + "\n".join(_bad_sentences) + "<br><br>"

            # time.sleep(2)

        self.statusLabel.setText(bad_sentences)


def run_application():
    "running application"
    app = qtw.QApplication(sys.argv)
    window = FetchDOIApp()  # pylint: disable=W0612
    sys.exit(app.exec_())
