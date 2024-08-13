# -*- coding: utf-8 -*-
"""
Created on Wed Aug  7 10:00:53 2024

@author: Ivan
"""
import os
import re
from functools import lru_cache
from typing import List, Dict, Optional, Tuple
from docx import Document
from docx.text.paragraph import Paragraph
import requests


# %% GLOBALS
LREF_PATTERN = re.compile(r'\[LR[^\]]*\]')

CN_BASE_URL = "https://doi.org"
HEADERS = {"User-Agent": "python-requests/" + requests.__version__,
           "X-USER-AGENT": "python-requests/" + requests.__version__,
           "Accept": "text/x-bibliography; style = apa; locale = en-Us"}

# %% FROM CT


@lru_cache(maxsize=256)
def fetch_request(url: str) -> Optional[str]:
    """
    Interface to fetch the request
    """
    r = requests.get(url, headers=HEADERS, allow_redirects=True)  # pylint: disable=C0103
    if not r.ok:
        return None
    r.encoding = "UTF-8"
    return r.text


def extract_references(docx_path: str,
                       pattern: Optional[re.Pattern] = None) -> Dict[str, List[str]]:
    """
    Extracts all the references found in
    a document path that match the given pattern.
    """
    if pattern is None:
        pattern = LREF_PATTERN

    doc = Document(docx_path)
    doi_list = list()
    for para in doc.paragraphs:
        matches = pattern.findall(para.text)
        doi_list.extend(matches)

    reference_dict = {k: k.strip("]").strip("[LR:").split(";") for k in doi_list}
    return reference_dict


def merge_dois(doi_dict: Dict[str, List[str]]) -> List[str]:
    """
    Merges all DOIs from the extracted references into a single list.
    """
    return list({item for sublist in doi_dict.values() for item in sublist})


def intext_cit(literature_citation: str) -> str:
    """
    Creates an APA style intext citation from a given string

    """
    authors, year = literature_citation.split("(")[:2]
    year = year.split(")")[0]

    author_count = authors.count(".,") + authors.count("&")

    first_author = authors.split(",")[0].strip()
    if author_count == 2:
        second_author = " & " + \
            authors.split("&")[1].split(",")[0].strip()+", "
    # ., counts all authors and & counts last author
    # when there is one author this can be only 0
    elif author_count == 0:
        first_author = first_author + ","
        second_author = " "
    else:
        second_author = " et al., "

    intext_citation = first_author + second_author + year
    return intext_citation


def generate_replacer(lrtext_doi: Dict[str, List[str]],
                      doi_intext: Dict[str, str]) -> Tuple[Dict[str, str], Dict[str, str]]:
    """
    "lrtext_doi" is how references are in the raw file, e.g.
    "doi_intext" is how a citation is written during text, e.g.
    '(Zhang et al., 2024)'

    This returns two dictionaries - good replacer and bad replacer
    """
    good_replacer = dict()
    bad_replacer = dict()
    for intext, doi_list in lrtext_doi.items():
        replacement = "("
        one_citation = len(doi_list) == 1

        if one_citation:
            joiner = ")"
            end = ""
        else:
            joiner = "; "
            end = ")"

        for doi in doi_list:
            intext_doi = doi_intext.get(doi, None)
            if intext_doi is None:
                intext_doi = "BAD_DOI:"+doi
            replacement = replacement + intext_doi + joiner

        replacement = replacement.rstrip("; ") + end

        if "BAD_DOI" in replacement:
            bad_replacer[intext] = replacement
        else:
            good_replacer[intext] = replacement
    return good_replacer, bad_replacer


def replace_text_in_runs(paragraph: Paragraph,
                         replacements: Dict[str, str]) -> None:
    """
    Replaces a pattern in text
    Doesnt save comments
    """
    for key, value in replacements.items():

        if key in paragraph.text:
            full_text = ''.join(run.text for run in paragraph.runs)
            full_text = full_text.replace(key, value)

            start_index = 0
            for run in paragraph.runs:
                run_length = len(run.text)
                run.text = full_text[start_index:start_index + run_length]
                start_index += run_length


def save_document(doc_path: str, doc_out: str, replacement: Dict[str, str]) -> None:
    """
    Replaces a pattern in doc and then saves it as the output
    """
    doc = Document(doc_path)
    for paragraph in doc.paragraphs:
        replace_text_in_runs(paragraph, replacement)
    doc.save(doc_out)


def add_suffix(file_path: str, suffix: str) -> str:
    """
    Add a suffix to the filename in the given file path.
    """
    directory, filename = os.path.split(file_path)
    name, ext = os.path.splitext(filename)

    new_filename = f"{name}{suffix}{ext}"

    # Construct the new file path
    new_file_path = os.path.join(directory, new_filename)

    return new_file_path


def highlight_text(text: str, pattern: str, color_code: str) -> str:
    """Highlight occurrences of `pattern` in `text` with the specified `color_code`."""
    parts = text.split(pattern)
    highlighted_text = ''

    for i in range(len(parts) - 1):
        highlighted_text += parts[i]

        highlighted_text += f'<font color="{color_code}">{pattern}</font>'

    highlighted_text += parts[-1]

    return highlighted_text


def split_into_sentences(text: str) -> List[str]:
    """Split text into sentences using basic punctuation rules."""

    sentence_endings = re.compile(r"(?<=[.!?])\s+(?=[A-Z])")
    sentences = sentence_endings.split(text)
    # edge cases
    return [sentence.strip() for sentence in sentences if sentence.strip()]


def recognize_bad_dois(doc_path: str, patterns: List[str]) -> List[str]:
    """
    Recognizes bad doi and returns sentences containing them,
    with highlighting them
    """
    doc = Document(doc_path)

    # I think orange is nice but red is even bettern
    highlight_color = "red"
    bad_sentences = list()

    for para in doc.paragraphs:
        text = para.text
        sentences = split_into_sentences(text)
        for sentence in sentences:
            if any(pattern in sentence for pattern in patterns):
                highlighted_sentence = sentence
                for pattern in patterns:
                    if pattern in highlighted_sentence:
                        highlighted_sentence = highlight_text(
                            highlighted_sentence, pattern, highlight_color)
                        bad_sentences.append(highlighted_sentence)
    return bad_sentences


def save_literature_to_docx(literature_entries: List[str], last_path: str) -> str:
    """
    saves literature to document provided
    """
    # literature saving loging
    end_literature = sorted(list(literature_entries))
    _directory, _ = os.path.split(last_path)  # pylint: disable=W0631
    _literature = os.path.join(_directory, "Literature.docx")

    doc = Document()
    doc.add_heading('Literature References', level=1)
    for entry in end_literature:
        doc.add_paragraph(entry.strip())  # Strip trailing newline characters
    doc.save(_literature)
    return _literature
