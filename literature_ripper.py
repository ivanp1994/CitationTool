# -*- coding: utf-8 -*-
"""
Created on Fri Aug  2 17:03:16 2024

@author: Ivan
"""
import os
import glob
from docx import Document
import re
from tqdm import tqdm
from habanero import cn
import pandas as pd

#%% extract references
def extract_references(docx_path:str):
    # Load the Word document
    doc = Document(docx_path)
    
    # Define the regex pattern to match [LREF: doilink]
    pattern = re.compile(r'\[LREF[^\]]*\]')
    
    # List to store all matched DOIs
    doi_list = []
    
    # Iterate through all paragraphs in the document
    for para in doc.paragraphs:
        # Find all matches in the paragraph text
        matches = pattern.findall(para.text)
        doi_list.extend(matches)
    
    return doi_list

def extract_references_multiple(docx_paths:list):
    doi_lists = list()
    for docx_path in docx_paths:
        doi_lists.extend(extract_references(docx_path))
    return list(set(doi_lists))

#%% habanero DOI
def extract_doi(doi):
    return cn.content_negotiation(ids = doi, format = "text", style = "apa")

#%% forming intext citation
def intext_cit(wupa:str):
    authors, year =  wupa.split("(")[:2]
    year = year.split(")")[0]
    
    author_count = authors.count(".,") + authors.count("&")
    
    first_author = authors.split(",")[0].strip()+" "
    if author_count == 2:
        second_author = "& "+ authors.split("&")[1].split(",")[0].strip()+", "
    elif author_count == 1:
        second_author = ""
    else:
        second_author = "et al., "
    
    intext_citation = first_author + second_author + year
    return intext_citation 

#%% MAIN FUNCTIONS
def create_literature_base(all_dois:list,force=False):
    if os.path.isfile("literature_buffer.csv") and not force:
        return pd.read_csv("literature_buffer.csv",index_col=0)
    
    for identif in all_dois:
        doi = identif.replace("[LREF:","").replace("]","").strip()
    
    good_dois = dict()
    bad_dois = dict()
    
    for identif in tqdm(all_dois, desc="Extracting DOIs"):
        doi = identif.replace("[LREF:","").replace("]","").strip()
        citation = extract_doi(doi)        
        if isinstance(citation,str):
            good_dois[identif]=citation
        elif isinstance(citation,list):
            for element in citation:
                if element is not None:
                    good_dois[identif]=element
                    break
            else:
                bad_dois[identif] = element
        else:
            bad_dois[identif] = element
            #continue
            #raise ValueError("Something went wrong with %s",k)
    
    good_df = pd.DataFrame([good_dois]).T.reset_index()
    good_df.columns = ["IDF","LITER"]
    good_df["INTEXT"] = good_df.LITER.apply(intext_cit)
    
    bad_df = pd.DataFrame([bad_dois]).T.reset_index()
    bad_df.columns = ["IDF","LITER"]
    bad_df["INTEXT"] = "-1"
    
    result_df = pd.concat([good_df,bad_df])
    result_df.to_csv("literature_buffer.csv")
    return create_literature_base(all_dois,False)

def save_good_literature(good_literature_df,folder="processed"):
    
    
    doc = Document()
    # Add each value to the document separated by two newlines
    for value in good_literature_df["LITER"].unique():
        doc.add_paragraph(value)
        #doc.add_paragraph("")  # Add an empty paragraph to create two newlines
    doc.save(f'{folder}/LITERATURE_FORMATTED.docx')
    return bad_literature

#%% replace files 
def replace_text_in_runs(paragraph, replacements):
    for key, value in replacements.items():
        if key in paragraph.text:
            # When a match is found, concatenate the text of all runs
            full_text = ''.join(run.text for run in paragraph.runs)
            # Replace the target text in the concatenated string
            full_text = full_text.replace(key, value)
            # Split the text back into runs
            start_index = 0
            for run in paragraph.runs:
                run_length = len(run.text)
                run.text = full_text[start_index:start_index + run_length]
                start_index += run_length

def save_document(doc_path,doc_out,replacement):
    doc = Document(doc_path)
    for paragraph in doc.paragraphs:
        replace_text_in_runs(paragraph, replacement)
    doc.save(doc_out)
    

files = ["4B_Rasprava_MMUS.docx","1_UVOD.docx","2_Materijali&Metode.docx","3A_Rezultati_AMEX.docx","3B_Rezultati_MMUS.docx","4A_Rasprava_AMEX.docx"]
all_dois = extract_references_multiple(files)
literature_base = create_literature_base(all_dois).sort_values("LITER")
bad_literature = literature_base.loc[literature_base.INTEXT=="-1"]
good_literature = literature_base.loc[literature_base.INTEXT!="-1"]

os.makedirs("processed",exist_ok=True)
save_good_literature(good_literature)


good_literature.INTEXT = "("+good_literature.INTEXT+")"
literature_dict = dict(zip(good_literature.IDF,good_literature.INTEXT)) 
for doc_part in files:
    doc_out = f"processed/{doc_part}"
    save_document(doc_part,doc_out,literature_dict)


