# -*- coding: utf-8 -*-
"""
Created on Wed Aug  7 10:00:53 2024

@author: Ivan
"""
import os
import re
import argparse
from functools import lru_cache
from typing import Dict, List, Optional, Tuple
import requests


from tqdm import tqdm
from docx import Document
from docx.text.paragraph import Paragraph


# %% GLOBALS
LREF_PATTERN = re.compile(r'\[LR[^\]]*\]')

CN_BASE_URL = "https://doi.org"
HEADERS = {"User-Agent": "python-requests/" + requests.__version__,
           "X-USER-AGENT": "python-requests/" + requests.__version__,
           "Accept": "text/x-bibliography; style = apa; locale = en-Us"}


# %% STARTING FUNCTIONS

def add_suffix(file_path: str, suffix: str) -> str:
    """
    Add a suffix to the filename in the given file path.
    """
    directory, filename = os.path.split(file_path)
    name, ext = os.path.splitext(filename)

    new_filename = f"{name}{suffix}{ext}"

    # Construct the new file path
    new_file_path = os.path.join(directory, new_filename)

    return new_file_path


def extract_references(docx_path: str,
                       pattern: Optional[re.Pattern] = None) -> Dict[str, List[str]]:
    """
    Extracts all the references
    found in a document path that match the given pattern. The given
    pattern is [LR:doi1;doi2]. A pattern can be changed by providing optional
    pattern argument or modifying LREF_PATTERN global variable.
    What is returned is a dictionary of intext pattern and List of strings
    """
    if pattern is None:
        pattern = LREF_PATTERN

    doc = Document(docx_path)

    doi_list = list()
    for para in doc.paragraphs:
        # Find all matches in the paragraph text
        matches = LREF_PATTERN.findall(para.text)
        doi_list.extend(matches)

    reference_dict = {k: k.strip("]").strip(
        "[LR:").split(";") for k in doi_list}
    return reference_dict


def green_text(text: str) -> str:
    """Return a string formatted with green color for terminal output."""
    green_color_code = "\033[92m"
    reset_color_code = "\033[0m"
    return f"{green_color_code}{text}{reset_color_code}"

# %% FUNCTIONS FOR RETRIEVING DOI


@lru_cache(maxsize=256)
def fetch_request(url: str) -> Optional[str]:
    """
    Interface to the request
    """
    r = requests.get(url, headers=HEADERS, allow_redirects=True,
                     )  # pylint: disable=C0103
    if not r.ok:
        return None
    r.encoding = "UTF-8"
    return r.text


def fetch_doi(doi_list: List[str],
              ) -> Dict[str, Optional[str]]:
    """
    Fetches an APA style format of the doi
    """
    doi_dictionary = dict()
    for ids in tqdm(doi_list, "Fetching dois"):
        doi_dictionary[ids] = fetch_request(CN_BASE_URL + "/" + ids.strip())
    return doi_dictionary

# %% FUNCTIONS FOR FORMATTING CITATIONS


def intext_cit(literature_citation: str) -> str:
    """
    Creates an APA style intext citation from a given string

    """
    authors, year = literature_citation.split("(")[:2]
    year = year.split(")")[0]

    author_count = authors.count(".,") + authors.count("&")

    first_author = authors.split(",")[0].strip()
    if author_count == 2:
        second_author = " & " + \
            authors.split("&")[1].split(",")[0].strip()+", "
    # ., counts all authors and & counts last author
    # when there is one author this can be only 0
    elif author_count == 0:
        first_author = first_author + ","
        second_author = " "
    else:
        second_author = " et al., "

    intext_citation = first_author + second_author + year
    return intext_citation


def generate_replacer(lrtext_doi: Dict[str, List[str]],
                      doi_intext: Dict[str, str]) -> Tuple[Dict[str, str], Dict[str, str]]:
    """
    "lrtext_doi" is how references are in the raw file, e.g.
    '[LR:  https://doi.org/10.1016/j.cmet.2024.07.004 ]' and the value is without LR
    "doi_intext" is how a citation is written during text, e.g.
    '(Zhang et al., 2024)'

    This returns two dictionaries - good replacer and bad replacer
    """
    good_replacer = dict()
    bad_replacer = dict()
    for intext, doi_list in lrtext_doi.items():
        replacement = "("
        one_citation = len(doi_list) == 1

        if one_citation:
            joiner = ")"
            end = ""
        else:
            joiner = "; "
            end = ")"

        for doi in doi_list:
            intext_doi = doi_intext.get(doi, None)
            if intext_doi is None:
                intext_doi = "BAD_DOI:"+doi
            replacement = replacement + intext_doi + joiner

        replacement = replacement.rstrip("; ") + end

        if "BAD_DOI" in replacement:
            bad_replacer[intext] = replacement
        else:
            good_replacer[intext] = replacement
    return good_replacer, bad_replacer

# %% BAD PATTERN RECOGNIZAL


def highlight_text(text: str, pattern: str, color_code: str) -> str:
    """Highlight occurrences of `pattern` in `text` with the specified `color_code`."""
    parts = text.split(pattern)
    highlighted_text = ''

    for i in range(len(parts) - 1):
        highlighted_text += parts[i]

        highlighted_text += f"{color_code}{pattern}\033[0m"

    highlighted_text += parts[-1]

    return highlighted_text


def split_into_sentences(text: str) -> List[str]:
    """Split text into sentences using basic punctuation rules."""

    sentence_endings = re.compile(r"(?<=[.!?])\s+(?=[A-Z])")
    sentences = sentence_endings.split(text)
    # edge cases
    return [sentence.strip() for sentence in sentences if sentence.strip()]


def recognize_bad_dois(doc_path: str, patterns: List[str]) -> List[str]:
    """
    Recognizes bad doi and returns sentences containing them,
    with highlighting them
    """
    doc = Document(doc_path)

    # I think orange is nice
    highlight_color = "\033[38;5;208m"
    bad_sentences = list()

    for para in doc.paragraphs:
        text = para.text
        sentences = split_into_sentences(text)
        for sentence in sentences:
            if any(pattern in sentence for pattern in patterns):
                highlighted_sentence = sentence
                for pattern in patterns:
                    if pattern in highlighted_sentence:
                        highlighted_sentence = highlight_text(
                            highlighted_sentence, pattern, highlight_color)
                        bad_sentences.append(highlighted_sentence)
    return bad_sentences

# %% END SAVING


def replace_text_in_runs(paragraph: Paragraph,
                         replacements: Dict[str, str]) -> None:
    """
    Replaces a pattern in text
    Doesnt save comments
    """
    for key, value in replacements.items():
        if key in paragraph.text:
            full_text = ''.join(run.text for run in paragraph.runs)
            full_text = full_text.replace(key, value)
            start_index = 0
            for run in paragraph.runs:
                run_length = len(run.text)
                run.text = full_text[start_index:start_index + run_length]
                start_index += run_length


def save_document(doc_path: str, doc_out: str, replacement: Dict[str, str]) -> None:
    """
    Replaces a pattern in doc and then saves it as the output
    """
    doc = Document(doc_path)
    for paragraph in doc.paragraphs:
        replace_text_in_runs(paragraph, replacement)
    doc.save(doc_out)


def save_literature_to_docx(literature_entries: List[str], doc_path: str):
    """
    saves literature to document provided
    """
    doc = Document()
    doc.add_heading('Literature References', level=1)
    for entry in literature_entries:
        doc.add_paragraph(entry.strip())  # Strip trailing newline characters
    doc.save(doc_path)
# %%  MAIN


def fix_citations(doc_paths: List[str],
                  pattern: Optional[re.Pattern] = None) -> None:
    """
    Main function that fixes citations.
    For every file in doc_path, an additional file with suffix '_proc'
    is created.
    The entire literature in APA format is found at the location of the
    last file as "Literature.docx"

    The pattern for doi recognizal is outlined in LREF_PATTERN global
    and it follows [LR:*] scheme with multiple citations for same sentence
    possible [LR:*;*]

    Optionally, the pattern can be changed via "pattern"


    """
    end_literature = set()
    if len(doc_paths) == 0:
        raise ValueError("No documents provided")
    for path in doc_paths:
        file_name = os.path.basename(path)
        result_name = add_suffix(path, "_proc")
        lrtext_doi = extract_references(
            path, pattern)  # converts it to lrtext_doi

        # there are all citations in this file
        merged_list = list({item for sublist in lrtext_doi.values()
                           for item in sublist})
        print("Done extracting from file ", green_text(file_name))
        doi_citation = fetch_doi(merged_list)  # <- the time consuming part

        print("Done fetching literature, the following bad dois are present in the file:")

        doi_intext = {k: intext_cit(v)
                      for k, v in doi_citation.items() if v is not None}
        end_literature = end_literature.union(
            {v for v in doi_citation.values() if v})
        good_replacer, bad_replacer = generate_replacer(lrtext_doi, doi_intext)
        bad_sentences = recognize_bad_dois(path, list(bad_replacer.keys()))
        for sentence in bad_sentences:
            print("\t", sentence)
        # replacing loging
        save_document(path, result_name, good_replacer)
        print("Saved file found at ", green_text(path))
        print("\n")

    # literature saving loging
    end_literature = sorted(list(end_literature))
    _directory, _ = os.path.split(path)  # pylint: disable=W0631
    _literature = os.path.join(_directory, "Literature.docx")
    save_literature_to_docx(end_literature, _literature)
    print("Saved literature to ", green_text(_literature))


# %%
if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Filling out references using their DOI")
    parser.add_argument("files", nargs='+', )
    files = vars(parser.parse_args())["files"]
    fix_citations(files)
